"""Recria os templates de O.S com tabelas limpas, mantendo as artes oficiais.

Base: layout medido nos PDFs originais (posições de textos e imagens via fitz).
Evita os artefatos de células mescladas do pdf2docx (texto duplicado,
sobreposição, linhas altas demais).

COMO USAR (manutenção dos modelos):
1. Rode este script:  python scripts/gerar_templates_os.py
   -> regenera backend/templates/OS_CONSTRUCAO.docx e OS_LINHA_VIVA.docx
2. Se quiser Ajustar o layout, abra os .docx gerados no Word e edite
   livremente — os marcadores {{campo}}, {% if %}, {%tr %} e {%p %} devem
   ser preservados (veja lista de marcadores abaixo).
3. As artes grandes vêm de arquivos PNG em cache local (IMGS). Se o cache
   não existir, extraia as imagens dos PDFs originais (manuais/MODELO O.S/)
   com:  python -c "from pdf2docx import Converter; Converter('...pdf').convert('...docx')"
   e depois extraia os blips do docx (rId9=logo, rId10=arte p1, rId11=arte p2,
   rId12=faixa LINHA VIVA).

MARCADORES DISPONÍVEIS:
  {{codigo}} {{equipe}} {{id_obra}} {{agencia}} {{data}} {{encarregado}}
  {{obra}} {{atividade}} {{local}} {{municipio}} {{bt}} {{at}}
  {{h_desligar}} {{h_religar}} {{alimentador}} {{chave}} {{servico}} {{obs}}
  membros (loop): {%tr for m in membros %} / {{ m.nome }} / {{ m.cargo }} / {%tr endfor %}
"""

import os

import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

DEST = r"C:\Users\User\Desktop\App_Munaretto\backend\templates"
IMGS = r"C:\Users\User\AppData\Local\Temp\opencode\imgs"
LOGO = os.path.join(IMGS, "orig_MODELO_CONSTRUÇÃO_rId9.png")  # mesmo logo nos dois

# Arquivos de arte extraídos dos docx convertidos
ARTES = {
    "construcao": {
        "p1": ("orig_MODELO_CONSTRUÇÃO_rId10.png", 18.1, 16.5),
        "p2": ("orig_MODELO_CONSTRUÇÃO_rId11.png", 18.1, 20.2),
    },
    "linha_viva": {
        "p1": ("orig_MODELO_LINHA_VIVA_rId10.png", 18.8, 17.1),
        "p2": ("orig_MODELO_LINHA_VIVA_rId11.png", 19.5, 17.5),
        "banner": ("orig_MODELO_LINHA_VIVA_rId12.png", 19.5, 2.0),
    },
}

FONTE = "Arial"


def _run(p, texto, negrito=False, tamanho=11, cor=None, central=False):
    r = p.add_run(texto)
    r.bold = negrito
    r.font.size = Pt(tamanho)
    r.font.name = FONTE
    if cor:
        r.font.color.rgb = cor
    return r


def _celula(cell, texto, negrito=False, tamanho=10, central=False, espaco=0):
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(espaco)
    p.paragraph_format.space_after = Pt(espaco)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if central else WD_ALIGN_PARAGRAPH.LEFT
    # remove runs existentes para não duplicar texto ao escrever 2x na célula
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    if texto:
        _run(p, texto, negrito=negrito, tamanho=tamanho)
    return p


def _altura_linha(tabela, indice, cm):
    tr = tabela.rows[indice]._tr
    tr_pr = tr.get_or_add_trPr()
    tr_h = OxmlElement("w:trHeight")
    tr_h.set(qn("w:val"), str(int(cm * 567)))  # 1cm = 567 twips
    tr_h.set(qn("w:hRule"), "exact")
    tr_pr.append(tr_h)


def _borda_celula(cell, espessura=4, cor="2F2F2F"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    bordas = OxmlElement("w:tcBorders")
    for lado in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{lado}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(espessura))
        el.set(qn("w:color"), cor)
        bordas.append(el)
    tc_pr.append(bordas)


def _tabela_grade(doc, cols_cm):
    """Cria tabela com bordas e colunas na largura dada (soma <= 18cm)."""
    tb = doc.add_table(rows=0, cols=len(cols_cm))
    tb.alignment = WD_TABLE_ALIGNMENT.CENTER
    # estilo: sem o Table Grid (bordas manuais por célula), com layout fixo
    tbl_pr = tb._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    # remove largura automática do grid (definimos por coluna)
    for grid_col, larg in zip(tb._tbl.tblGrid.findall(qn("w:gridCol")), cols_cm, strict=True):
        grid_col.set(qn("w:w"), str(int(larg * 567)))
    return tb


def _add_linha(tb, valores, largura_cm, tamanho=10, negrito=False, altura=None):
    linha = tb.add_row()
    for cell, texto in zip(linha.cells, valores, strict=True):
        _borda_celula(cell)
        _celula(cell, texto, negrito=negrito, tamanho=tamanho)
    if altura:
        _altura_linha(tb, len(tb.rows) - 1, altura)
    return linha


def _imagem(doc, caminho, largura_cm, altura_cm, central=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if central else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    run.add_picture(caminho, width=Cm(largura_cm), height=Cm(altura_cm))
    return p


def _texto(doc, texto, negrito=False, tamanho=11, central=False, antes=0, depois=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if central else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(antes)
    p.paragraph_format.space_after = Pt(depois)
    if texto:
        _run(p, texto, negrito=negrito, tamanho=tamanho)
    return p


def _texto_compacto(doc, texto, tamanho=10):
    """Parágrafo com altura de linha mínima (para os marcadores {%p %} do loop)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(12)  # linha exata ~0.42cm
    if texto:
        _run(p, texto, tamanho=tamanho)
    return p


def criar(tipo, titulo):
    doc = docx.Document()
    secao = doc.sections[0]
    secao.top_margin = Cm(1.2)
    secao.bottom_margin = Cm(1.2)
    secao.left_margin = Cm(1.5)
    secao.right_margin = Cm(1.5)
    secao.page_width = Cm(21)
    secao.page_height = Cm(29.7)

    artes = ARTES[tipo]

    # ===== Página 1 =====
    # Tabela principal: logo + título + campos (larguras: rótulos não quebram)
    cols = [2.6, 3.2, 2.6, 2.2, 2.2, 2.6, 1.3, 1.3]  # total 18cm
    tb = _tabela_grade(doc, cols)

    # R0: logo + título (span 5) + vazios
    r0 = tb.add_row()
    _borda_celula(r0.cells[0])
    _celula(r0.cells[0], "")
    r0.cells[0].paragraphs[0].add_run().add_picture(LOGO, width=Cm(2.3), height=Cm(1.5))
    # mescla células 1-5 para o título
    cel_titulo = r0.cells[1]
    for c in r0.cells[2:6]:
        cel_titulo.merge(c)
    _borda_celula(cel_titulo)
    _celula(cel_titulo, titulo, negrito=True, tamanho=16, central=True)
    for c in r0.cells[6:]:
        _borda_celula(c)
        _celula(c, "")

    # R1: código / equipe / id obra / agência
    _add_linha(
        tb,
        ["Ord.Serv. {{codigo}}", "", "Equipe", "{{equipe}}", "ID Obra", "{{id_obra}}", "Agência", "{{agencia}}"],
        cols,
        tamanho=10,
    )
    # mescla as 2 primeiras células do R1 (código)
    linha1 = tb.rows[-1]
    cel_cod = linha1.cells[0].merge(linha1.cells[1])
    _celula(cel_cod, "Ord.Serv. {{codigo}}", negrito=True, tamanho=10)

    # R2: Data / Encarregado
    _add_linha(tb, ["Data", "{{data}}", "Encarregado", "{{encarregado}}", "", "", "", ""], cols, tamanho=10)
    linha2 = tb.rows[-1]
    cel_enc = linha2.cells[3].merge(linha2.cells[4])
    _celula(cel_enc, "{{encarregado}}", tamanho=10)

    # R3: Obra / Atividade
    _add_linha(tb, ["Obra", "{{obra}}", "Atividade", "", "", "{{atividade}}", "", ""], cols, tamanho=10)
    linha3 = tb.rows[-1]
    cel_ativ = linha3.cells[2].merge(linha3.cells[4])
    _celula(cel_ativ, "Atividade", negrito=True, tamanho=10)

    # R4: Local / Município
    _add_linha(tb, ["Local", "{{local}}", "Município", "", "", "{{municipio}}", "", ""], cols, tamanho=10)
    linha4 = tb.rows[-1]
    cel_mun = linha4.cells[2].merge(linha4.cells[4])
    _celula(cel_mun, "Município", negrito=True, tamanho=10)

    # Tabela BT/AT + horários + alimentador/chave
    cols_b = [2.4, 1.6, 2.4, 1.0, 2.8, 2.2, 2.8, 2.8]  # 18cm
    tb_b = _tabela_grade(doc, cols_b)
    _add_linha(
        tb_b,
        [
            "BT Energ.",
            "{% if bt %}X{% endif %}",
            "",
            "",
            "H.Desligar",
            "{{h_desligar}}",
            "Alimentador",
            "{{alimentador}}",
        ],
        cols_b,
        tamanho=9,
    )
    _add_linha(
        tb_b,
        ["AT Energ.", "{% if at %}X{% endif %}", "Bloqueio", "", "H.Religar", "{{h_religar}}", "Chave", "{{chave}}"],
        cols_b,
        tamanho=9,
    )

    # Serviço a Executar + Obs. (com os valores)
    _texto(doc, "Serviço a Executar", negrito=True, tamanho=10, antes=6)
    _texto(doc, "{{servico}}", tamanho=10, depois=2)
    _texto(doc, "Obs.", negrito=True, tamanho=10, antes=4)
    _texto(doc, "{{obs}}", tamanho=10, depois=4)

    # Arte da página 1 (embaixo, como no original)
    img1, larg1, alt1 = artes["p1"]
    _imagem(doc, os.path.join(IMGS, img1), larg1, alt1)

    # ===== Página 2 =====
    # Arte da página 2 (topo) — quebra de página via page_break_before
    img2, larg2, alt2 = artes["p2"]
    p_art2 = _imagem(doc, os.path.join(IMGS, img2), larg2, alt2)
    p_art2.paragraph_format.page_break_before = True
    if "banner" in artes:
        img3, larg3, alt3 = artes["banner"]
        _imagem(doc, os.path.join(IMGS, img3), larg3, alt3)

    # Título da avaliação prévia + data
    tb_c = _tabela_grade(doc, [2.5, 13.0, 2.5])
    r = tb_c.add_row()
    _borda_celula(r.cells[0])
    _celula(r.cells[0], "")
    r.cells[0].paragraphs[0].add_run().add_picture(LOGO, width=Cm(1.4), height=Cm(0.95))
    _borda_celula(r.cells[1])
    _celula(
        r.cells[1], "AVALIAÇÃO PRÉVIA DE ESTUDO E PLANEJAMENTO DAS ATIVIDADES", negrito=True, tamanho=11, central=True
    )
    _borda_celula(r.cells[2])
    _celula(r.cells[2], "{{data}}", tamanho=10, central=True)

    if tipo == "construcao":
        # Tabela de membros
        tb_m = _tabela_grade(doc, [7.0, 5.0, 6.0])
        _add_linha(tb_m, ["Membros da Equipe", "", "O.S. Nr. : {{codigo}}"], [7.0, 5.0, 6.0], negrito=True, tamanho=10)
        linha_m0 = tb_m.rows[-1]
        cel = linha_m0.cells[0].merge(linha_m0.cells[1])
        _celula(cel, "Membros da Equipe", negrito=True, tamanho=10)
        _add_linha(tb_m, ["Nome", "Cargo", "Assinatura"], [7.0, 5.0, 6.0], negrito=True, tamanho=10)
        # loop
        _add_linha(tb_m, ["{%tr for m in membros %}", "", ""], [7.0, 5.0, 6.0], tamanho=10, altura=0.2)
        _add_linha(tb_m, ["{{ m.nome }}", "{{ m.cargo }}", ""], [7.0, 5.0, 6.0], tamanho=10, altura=0.58)
        _add_linha(tb_m, ["{%tr endfor %}", "", ""], [7.0, 5.0, 6.0], tamanho=10, altura=0.2)
    else:
        # LINHA VIVA: O.S. Nr. + membros em parágrafos compactos
        _texto(doc, "O.S. Nr. : {{codigo}}", negrito=True, tamanho=10, antes=4)
        _texto(doc, "Membros da Equipe", negrito=True, tamanho=10, antes=2)
        _texto(doc, "Nome\tCargo\tAssinatura", negrito=True, tamanho=10, antes=2)
        _texto_compacto(doc, "{%p for m in membros %}")
        _texto_compacto(doc, "{{ m.nome }}\t{{ m.cargo }}\t____________________")
        _texto_compacto(doc, "{%p endfor %}")
        _texto(
            doc,
            "Data : ______/______/__________     Hora : _____/_____     "
            "Todos entenderam os requisitos de segurança? ____",
            tamanho=10,
            antes=4,
        )

    caminho = os.path.join(DEST, f"OS_{tipo.upper()}.docx")
    doc.save(caminho)
    print("Gerado:", caminho)


criar("construcao", "Ordem de Serviço")
criar("linha_viva", "Ordem de Serviço - Linha Viva")
