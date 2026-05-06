import os
import shutil
import webbrowser
from datetime import datetime

TEMPLATES_DIR = "templates"

# Templates padrão vazios - todos devem ser importados como arquivos .docx
TEMPLATES = {}


def ensure_templates_dir():
    """Cria a pasta de templates de documento se não existir."""
    os.makedirs(TEMPLATES_DIR, exist_ok=True)


def carregar_templates_externos():
    """Carrega templates em Word (.docx) da pasta templates."""
    ensure_templates_dir()
    templates = {}
    try:
        from docx import Document
        for nome_arquivo in sorted(os.listdir(TEMPLATES_DIR)):
            if nome_arquivo.lower().endswith(".docx"):
                chave = os.path.splitext(nome_arquivo)[0]
                caminho = os.path.join(TEMPLATES_DIR, nome_arquivo)
                try:
                    doc = Document(caminho)
                    conteudo = "\n".join([paragrafo.text for paragrafo in doc.paragraphs])
                    templates[chave] = {
                        "nome": f"Modelo: {chave}",
                        "template": conteudo,
                        "docx_path": caminho,
                    }
                except Exception as e:
                    print(f"Erro ao carregar template {nome_arquivo}: {e}")
    except ImportError:
        print("python-docx não instalado. Templates Word não carregados.")
    return templates


def get_templates():
    """Retorna todos os templates disponíveis, incluindo modelos carregados."""
    templates = dict(TEMPLATES)
    templates.update(carregar_templates_externos())
    return templates


def get_template(tipo_documento):
    """Retorna o template solicitado, incluindo modelos externos."""
    return get_templates().get(tipo_documento)


def importar_template_arquivo(caminho_origem):
    """Importa um arquivo Word (.docx) para a pasta de templates."""
    ensure_templates_dir()
    if not os.path.exists(caminho_origem):
        return False, "Arquivo não encontrado."
    if not caminho_origem.lower().endswith(".docx"):
        return False, "Apenas arquivos .docx (Word) são suportados como modelo."

    nome_arquivo = os.path.basename(caminho_origem)
    destino = os.path.join(TEMPLATES_DIR, nome_arquivo)
    try:
        shutil.copyfile(caminho_origem, destino)
        return True, f"Template importado com sucesso: {destino}"
    except Exception as e:
        return False, f"Erro ao importar template: {str(e)}"


def criar_pasta_documentos():
    """Garante que a pasta de documentos gerados exista."""
    if not os.path.exists("documentos_gerados"):
        os.makedirs("documentos_gerados")


def gerar_documento_txt(cliente_info, tipo_documento):
    """Gera documento em formato TXT."""
    criar_pasta_documentos()

    template_info = get_template(tipo_documento)
    if not template_info:
        return None

    data_atual = datetime.now().strftime("%d/%m/%Y")
    template = template_info["template"]

    conteudo = template.format(
        nome=cliente_info[1],
        cpf_cnpj=cliente_info[2],
        endereco=cliente_info[3],
        data=data_atual,
    )

    nome_arquivo = f"{cliente_info[1].replace(' ', '_')}_{tipo_documento}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
    caminho = os.path.join("documentos_gerados", nome_arquivo)

    with open(caminho, "w", encoding="utf-8") as f:
        f.write(conteudo)

    return caminho


def gerar_documento_html(cliente_info, tipo_documento):
    """Gera documento em formato HTML para visualização no navegador."""
    criar_pasta_documentos()

    template_info = get_template(tipo_documento)
    if not template_info:
        return None

    data_atual = datetime.now().strftime("%d/%m/%Y")
    template = template_info["template"]

    conteudo = template.format(
        nome=cliente_info[1],
        cpf_cnpj=cliente_info[2],
        endereco=cliente_info[3],
        data=data_atual,
    )

    html_content = f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{template_info['nome']}</title>
    <style>
        body {{
            font-family: 'Arial', sans-serif;
            max-width: 900px;
            margin: 20px auto;
            padding: 20px;
            background-color: #f5f5f5;
        }}
        .container {{
            background-color: white;
            padding: 40px;
            border-radius: 5px;
            box-shadow: 0 2px 5px rgba(0,0,0,0.1);
        }}
        h1 {{
            text-align: center;
            color: #333;
            margin-bottom: 30px;
        }}
        .content {{
            line-height: 1.8;
            color: #555;
            white-space: pre-wrap;
            word-wrap: break-word;
        }}
        .print-button {{
            text-align: center;
            margin-top: 30px;
        }}
        button {{
            background-color: #4CAF50;
            color: white;
            padding: 10px 30px;
            border: none;
            border-radius: 4px;
            cursor: pointer;
            font-size: 16px;
        }}
        button:hover {{
            background-color: #45a049;
        }}
        @media print {{
            body {{ background-color: white; }}
            .print-button {{ display: none; }}
        }}
    </style>
</head>
<body>
    <div class="container">
        <h1>{template_info['nome']}</h1>
        <div class="content">
{conteudo}
        </div>
        <div class="print-button">
            <button onclick="window.print()">🖨️ Imprimir</button>
        </div>
    </div>
</body>
</html>
"""

    nome_arquivo = f"{cliente_info[1].replace(' ', '_')}_{tipo_documento}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html"
    caminho = os.path.join("documentos_gerados", nome_arquivo)

    with open(caminho, "w", encoding="utf-8") as f:
        f.write(html_content)

    return caminho


def gerar_documento_word(cliente_info, tipo_documento):
    """Gera documento em formato Word (.docx)."""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        criar_pasta_documentos()

        template_info = get_template(tipo_documento)
        if not template_info:
            return None

        data_atual = datetime.now().strftime("%d/%m/%Y")
        template_text = template_info["template"]

        conteudo = template_text.format(
            nome=cliente_info[1],
            cpf_cnpj=cliente_info[2],
            endereco=cliente_info[3],
            data=data_atual,
        )

        # Se há um arquivo .docx como base, usar como template
        if "docx_path" in template_info and os.path.exists(template_info["docx_path"]):
            doc = Document(template_info["docx_path"])
            # Substituir placeholders nos parágrafos existentes
            for paragrafo in doc.paragraphs:
                if "{nome}" in paragrafo.text or "{cpf_cnpj}" in paragrafo.text or "{endereco}" in paragrafo.text or "{data}" in paragrafo.text:
                    paragrafo.text = paragrafo.text.format(
                        nome=cliente_info[1],
                        cpf_cnpj=cliente_info[2],
                        endereco=cliente_info[3],
                        data=data_atual,
                    )
        else:
            # Criar documento do zero
            doc = Document()
            titulo = doc.add_heading(template_info["nome"], level=1)
            titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for paragrafo_texto in conteudo.split("\n"):
                if paragrafo_texto.strip():
                    p = doc.add_paragraph(paragrafo_texto)
                    p.paragraph_format.line_spacing = 1.5

        nome_arquivo = f"{cliente_info[1].replace(' ', '_')}_{tipo_documento}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        caminho = os.path.join("documentos_gerados", nome_arquivo)
        doc.save(caminho)

        return caminho
    except ImportError:
        return None


def gerar_documento_pdf(cliente_info, tipo_documento):
    """Gera documento em formato PDF (.pdf)."""
    try:
        from fpdf import FPDF

        criar_pasta_documentos()

        template_info = get_template(tipo_documento)
        if not template_info:
            return None

        data_atual = datetime.now().strftime("%d/%m/%Y")
        template = template_info["template"]

        conteudo = template.format(
            nome=cliente_info[1],
            cpf_cnpj=cliente_info[2],
            endereco=cliente_info[3],
            data=data_atual,
        )

        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", "B", 16)
        pdf.cell(0, 10, template_info["nome"], ln=True, align="C")
        pdf.ln(5)
        
        pdf.set_font("Arial", "", 11)
        pdf.multi_cell(0, 5, conteudo)

        nome_arquivo = f"{cliente_info[1].replace(' ', '_')}_{tipo_documento}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        caminho = os.path.join("documentos_gerados", nome_arquivo)

        pdf.output(caminho)

        return caminho
    except ImportError:
        return None


def abrir_no_navegador(caminho_arquivo):
    """Abre arquivo no navegador padrão (funciona para HTML e PDFs)."""
    caminho_abs = os.path.abspath(caminho_arquivo)
    if os.path.exists(caminho_abs):
        webbrowser.open("file://" + caminho_abs)
        return True
    return False
